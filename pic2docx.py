from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.document import Document as Doc
import os
import sys
import datetime
from docx2pdf import convert
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class gen_docx():

    def __init__(self, pic_paths, extra_text) -> None:
        self.pic_paths = pic_paths
        self.extra_texts = extra_text
        self.get_pdf_path()
        self.create_document()

    def create_document(self):
        self.document = Document()  # 创建文档实例
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

    def get_pdf_path(self):
        current_time = datetime.datetime.now()
        formatted_time = current_time.strftime('%Y-%m-%d_%H-%M-%S')
        self.save_pdf_path = os.path.join(os.getcwd(), formatted_time + '.pdf')
        self.save_docx_path = os.path.join(os.getcwd(),
                                           formatted_time + '.docx')

    def put_pic_in_word(self):
        num_batch = len(self.pic_paths) // 18 + 1
        for batch_idx in range(num_batch):
            table = self.document.add_table(rows=9, cols=2)
            for i, (pic_path, extra_text) in enumerate(
                    zip(self.pic_paths[batch_idx * 18:(batch_idx + 1) * 18],
                        self.extra_texts[batch_idx * 18:(batch_idx + 1) *
                                         18])):
                # if i + 1 <= 9:
                #     cur_cell = self.table.cell(i, 0)
                #     cur_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中
                #     cur_cell.paragraphs[
                #         0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 水平居中
                #     run = cur_cell.paragraphs[0].add_run()  #添加图片的行
                #     run.add_picture(pic_path, height=Cm(2.25))
                # else:
                #     cur_cell = self.table.cell((i + 1) % 9 - 1, 1)
                #     cur_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直居中
                #     cur_cell.paragraphs[
                #         0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 水平居中
                #     run = cur_cell.paragraphs[0].add_run()  #添加图片的行
                #     run.add_picture(pic_path, height=Cm(2.25))
                cur_cell = table.cell(int(i / 2), i % 2)
                
                # tc = cur_cell._tc
                # # 创建一个新的边框元素
                # borders = parse_xml(r'<w:tcBorders %s><w:top w:val="dashed" w:sz="4"/><w:left w:val="dashed" w:sz="4"/><w:bottom w:val="dashed" w:sz="4"/><w:right w:val="dashed" w:sz="4"/></w:tcBorders>' % nsdecls('w'))
                # # 设置单元格的边框
                # tc.get_or_add_tcPr().append(borders)
                
                # 假设 cur_cell 是你当前操作的单元格实例
                cur_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格内容垂直居中
                # 设置单元格中第一个段落的水平对齐方式为居中
                cur_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = cur_cell.paragraphs[0].add_run()  # 添加图片的行
                run.add_picture(pic_path, height=Cm(2.25))
                if extra_text != '':
                    paragraph = cur_cell.add_paragraph()
                    paragraph.add_run(extra_text)
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置文字居中
                    paragraph_format.line_spacing = Pt(0)
                    paragraph_format.space_before = Pt(0)  # 设置段前距离为0
        self.document.save(self.save_docx_path)
        convert(self.save_docx_path)
        for pic_path in self.pic_paths:
            os.remove(pic_path)
        os.remove(self.save_docx_path)
